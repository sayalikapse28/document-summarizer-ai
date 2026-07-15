"""
document_parser.py
-------------------
Extracts raw text from common document formats: PDF, DOCX, XLSX/XLS, CSV, TXT.
Each file type has a dedicated extractor so we can handle its structure
(pages, paragraphs, tables, sheets) correctly.
"""

from pathlib import Path
import pandas as pd
from pypdf import PdfReader
from docx import Document


def extract_text(uploaded_file) -> str:
    """
    Detects file type from the filename extension and routes to the
    correct extractor. `uploaded_file` is a Streamlit UploadedFile
    (file-like object) or any object with .name and .read()/seek support.
    """
    name = uploaded_file.name
    ext = Path(name).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(uploaded_file)
    elif ext == ".docx":
        return _extract_docx(uploaded_file)
    elif ext in (".xlsx", ".xls"):
        return _extract_excel(uploaded_file)
    elif ext == ".csv":
        return _extract_csv(uploaded_file)
    elif ext == ".txt":
        return uploaded_file.read().decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file) -> str:
    reader = PdfReader(file)
    text_parts = []
    for page_num, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_docx(file) -> str:
    doc = Document(file)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of tables (common in reports/resumes/invoices)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


def _extract_excel(file) -> str:
    xls = pd.ExcelFile(file)
    parts = []
    for sheet_name in xls.sheet_names:
        df = xls.parse(sheet_name)
        parts.append(f"--- Sheet: {sheet_name} ---")
        parts.append(df.to_string(index=False))
    return "\n".join(parts)


def _extract_csv(file) -> str:
    df = pd.read_csv(file)
    return df.to_string(index=False)
